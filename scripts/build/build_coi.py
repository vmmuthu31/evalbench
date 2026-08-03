"""
Build the Elsevier-style Declaration of Competing Interest as a .docx file.
Run:  python3 evalbench_paper2/paper_draft/build_coi.py
"""

from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT     = Path(__file__).parent.parent.parent
OUT_PATH = ROOT / "evalbench_paper2/paper_draft/EvalBench_DeclarationOfCompetingInterest.docx"

doc = Document()

section = doc.sections[0]
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin   = Inches(1.2)
section.right_margin  = Inches(1.2)


def body(text, indent=False, bold=False, center=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(10)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold
    return p


body("Declaration of Competing Interest", bold=True, center=True)
doc.add_paragraph()

body(
    "Manuscript title: EvalBench v1 (Pilot Release): A Unified Multi-Task Benchmark for AI "
    "Evaluation in Educational Settings"
)
body("Authors: Vairamuthu M., Sudhan M. B.")
doc.add_paragraph()

body(
    "The authors declare that they have no known competing financial interests or personal "
    "relationships that could have appeared to influence the work reported in this paper.",
    indent=True,
)

body(
    "Specifically, the authors confirm that: (1) none of the authors have any financial interest "
    "or benefit arising from the direct applications of this research; (2) none of the authors "
    "are currently serving, or have previously served, in an editorial capacity for Engineering "
    "Applications of Artificial Intelligence; and (3) there are no other relationships or "
    "activities that readers could perceive to have influenced, or that give the appearance of "
    "potentially influencing, the work reported in this paper.",
    indent=True,
)

doc.add_paragraph()
body("Vairamuthu M.")
body("Sudhan M. B.")
body("August 3, 2026")

doc.save(str(OUT_PATH))
print(f"Saved: {OUT_PATH}")
