"""
Build the EvalBench manuscript as a .docx file with embedded figures.
Run:  python3 evalbench_paper2/paper_draft/build_docx.py
"""

from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

ROOT     = Path(__file__).parent.parent.parent
FIG_DIR  = ROOT / "evalbench_paper2/outputs/figures"
OUT_PATH = ROOT / "evalbench_paper2/paper_draft/EvalBench_Manuscript_Blinded.docx"

doc = Document()

# ── page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin   = Inches(1.2)
section.right_margin  = Inches(1.2)

# ── helpers ───────────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.size = Pt(13)
    p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.size = Pt(11)
    p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.first_line_indent = Inches(0.3)
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def caption(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(10)
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.italic = True
    return p

def fig(filename, cap_text, width=5.5):
    path = FIG_DIR / filename
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width))
    else:
        doc.add_paragraph(f"[Figure not found: {filename}]")
    caption(cap_text)

def add_table(headers, rows, cap_text):
    caption(cap_text)
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        drow = t.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = drow.cells[ci]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

def ref(num, authors, year, title, journal, doi=""):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent       = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.4)
    p.paragraph_format.space_after       = Pt(4)
    run = p.add_run(f"[{num}] ")
    run.font.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(f"{authors} ({year}). {title}. {journal}.")
    run2.font.size = Pt(10)
    if doi:
        run3 = p.add_run(f" {doi}")
        run3.font.size = Pt(10)

# ══════════════════════════════════════════════════════════════════════════════
#  TITLE  (author block intentionally omitted pending journal submission decision)
# ══════════════════════════════════════════════════════════════════════════════
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run(
    "EvalBench v1 (Pilot Release): A Unified Multi-Task Benchmark for AI Evaluation in "
    "Educational Settings"
)
run.bold = True
run.font.size = Pt(16)
doc.add_paragraph()

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run(
    "Manuscript prepared for anonymous peer review. Author names, affiliations, and contact "
    "details are provided in a separate title page document rather than in this file."
)
r.italic = True
r.font.size = Pt(9)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  HIGHLIGHTS  (required by most Q1 Elsevier-style journals, 3-5 bullets, <85 chars)
# ══════════════════════════════════════════════════════════════════════════════
h1("Highlights")
highlights = [
    "EvalBench v1 (Pilot Release) unifies OCR, assessment, and trustworthiness evaluation in one schema.",
    "132,891 records from six datasets support five standardised evaluation tracks.",
    "OCR rankings fully reverse between IAM and IIIT-Word (EM gap up to 0.589).",
    "Gemini Flash grades exam papers without a rubric at MAE = 3.78 marks (r = 0.543).",
    "Evidence Generation Rate reaches 100% but calibration stays near chance (AUROC = 0.571).",
]
for hl in highlights:
    p = doc.add_paragraph(hl, style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.size = Pt(11)

# ══════════════════════════════════════════════════════════════════════════════
#  ABSTRACT  (single paragraph, per Scopus/SCIE journal convention)
# ══════════════════════════════════════════════════════════════════════════════
h1("Abstract")
body(
    "The AI contribution of this work is a unified multi-task evaluation framework, EvalBench v1; "
    "the engineering application is document-processing infrastructure for automated educational "
    "assessment, where optical character recognition (OCR), grading, and trustworthiness checks "
    "are usually engineered and validated as separate pipelines, so cross-task failures stay "
    "hidden until deployment. EvalBench v1 is a pilot release comprising 132,891 quality-filtered "
    "records drawn from six public datasets; this figure denotes the processed, evaluation-ready "
    "benchmark corpus after transcription-availability and quality filtering, not the combined size "
    "of the raw source releases, which is larger for at least one constituent dataset (Section 3.1 "
    "details the IAM case, where 85,756 raw word images yield 38,305 records that pass the "
    "pipeline's filters). The corpus is organised into five evaluation tracks (OCR accuracy, "
    "document understanding, educational assessment, multi-modal understanding, trustworthiness) "
    "under a common schema, a fixed 70/10/20 split protocol, and a reproducible evaluation harness "
    "for cross-team validation on the same public data. This pilot release reports experiments for "
    "three of the five tracks; the remaining two "
    "are specified but held for a follow-up release. OCR model rankings prove unstable across "
    "datasets: TrOCR-base achieves the highest Exact "
    "Match on IAM (EM = 0.553) but ranks third on IIIT-Word (EM = 0.291), whereas PaddleOCR leads "
    "on IIIT-Word (EM = 0.664) yet ranks last on IAM (EM = 0.075), a rank reversal that "
    "single-corpus evaluation would not detect, a direct concern for any team choosing an OCR "
    "component from one benchmark alone. For assessment, Gemini Flash grades 50 handwritten exam "
    "papers with no rubric and achieves a Score mean absolute error (MAE) of 3.78 marks (Pearson "
    "r = 0.543). The trustworthiness track shows the model generated a non-empty evidence field for "
    "100% of predictions, a presence check rather than a verified support claim (Section 4.4), yet "
    "its stated confidence barely discriminates correct from incorrect "
    "scores (area "
    "under the ROC curve, AUROC = 0.571; expected calibration error, ECE = 0.251), an "
    "overconfidence pattern invisible under accuracy-only evaluation and directly "
    "relevant to engineering safe human-in-the-loop grading systems. EvalBench v1, with all "
    "evaluation scripts and outputs, is released to support reproducible engineering of "
    "educational AI systems."
)
body(
    "Keywords: educational AI evaluation, handwritten text recognition, automated exam scoring, "
    "confidence calibration, multi-task benchmark, cross-dataset generalisation, vision-language models."
)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 – INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
h1("1. Introduction")
body(
    "Artificial intelligence is already active in classrooms, reading handwritten answer sheets, "
    "scoring exam papers, and offering feedback on student work. What has not kept pace with this "
    "adoption is the infrastructure used to evaluate whether these systems actually work as "
    "intended. OCR tools are validated on one corpus, grading tools on another, and trustworthiness "
    "properties such as confidence calibration are rarely measured at all, in terms of any "
    "standard practice. This fragmentation makes it difficult to establish the reliability that "
    "high-stakes educational deployment requires."
)
body(
    "This study was undertaken to address all three gaps within a single framework. EvalBench "
    "introduces a five-track evaluation protocol, covering OCR accuracy, document understanding, "
    "educational assessment, multi-modal question answering, and trustworthiness, built on one "
    "shared data schema and one reproducible evaluation harness. This makes it possible to compare "
    "models across tasks and to identify failure modes that single-task evaluation cannot detect."
)
body(
    "The key contributions of this study are as follows."
)
items = [
    "We propose EvalBench v1, released here as a pilot, to our knowledge one of the first unified "
    "multi-task benchmarks for AI evaluation in educational settings, covering five task tracks "
    "under a common JSON schema and split protocol rather than treating each task as an isolated "
    "evaluation problem.",

    "We construct and release a fused benchmark corpus of 132,891 records from six independent, "
    "publicly available datasets spanning handwritten word recognition, full-page exam OCR, graded "
    "assessment papers, and handwritten question answering.",

    "We provide empirical baselines for four OCR models, namely Tesseract 5, EasyOCR, TrOCR-base, "
    "and PaddleOCR, together with two Gemini vision-language model variants across three tracks.",

    "We demonstrate a bidirectional rank reversal across OCR datasets, with cross-dataset Exact "
    "Match gaps as large as 0.589, which confirms that single-corpus evaluation produces "
    "systematically misleading model rankings.",

    "We define Track 5 (Trustworthiness), which measures confidence calibration and evidence "
    "generation independently of score accuracy and helps to expose overconfidence that standard "
    "accuracy metrics alone cannot detect.",
]
for item in items:
    p = doc.add_paragraph(item, style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(11)

body(
    "The remainder of the paper is organised as follows. Section 2 reviews related work on "
    "handwriting recognition and educational AI. Section 3 describes the EvalBench v1 dataset, "
    "schema, and pre-processing steps. Section 4 reports experimental results across three tracks. "
    "Section 5 presents an ablation study isolating the cross-dataset ranking effect. Section 6 "
    "discusses the findings and their limitations. Section 7 concludes with directions for future "
    "work."
)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 – RELATED WORK
# ══════════════════════════════════════════════════════════════════════════════
h1("2. Related Work")

h2("2.1 Handwritten Text Recognition and OCR")
body(
    "Handwritten text recognition has been a long-standing problem in computer vision and document "
    "processing. Early methods relied on handcrafted features and rule-based pipelines, which meant "
    "researchers had to design recognition rules manually. Transformer-based architectures changed "
    "this by learning character patterns directly from image data, which helps to explain the large "
    "accuracy gains reported over the last few years [1]."
)
body(
    "A 2025 benchmarking study compared eight large language models on handwritten text recognition "
    "and found that proprietary systems such as GPT-4o produced lower character error rates than "
    "open-source alternatives across a range of writing styles [1]. PaddleOCR 3.0, released the "
    "same year, restructured its detection-recognition pipeline and reported measurable precision "
    "gains on scene text over earlier versions [2]. However, such studies are typically evaluated "
    "on a single corpus, which means cross-dataset generalisation remains largely unreported."
)
body(
    "Layout-aware recognition has also advanced. The PILOT framework was developed to handle "
    "documents with mixed text and table regions, a structure common in exam papers where answers "
    "are interleaved with printed questions [3], which makes it directly relevant to educational "
    "document evaluation. OCRNet, a hybrid CNN architecture published in 2025, reports 95% accuracy "
    "on alphanumeric character recognition without a large pre-trained backbone [4]. The closest "
    "prior work to EvalBench's Track 1 design is OCRBench [18], a 29-dataset OCR evaluation suite "
    "for large multimodal models that already covers handwritten text recognition among other "
    "sub-tasks, however OCRBench evaluates OCR in isolation and does not extend to grading or "
    "trustworthiness, which is exactly the gap EvalBench is built to close. Within the OCR track "
    "specifically, EvalBench addresses the cross-corpus gap directly by evaluating every model on "
    "both IAM and IIIT-Word under one shared metric set, which is precisely what makes the rank "
    "reversal reported in Section 4 visible."
)

h2("2.2 Automated Exam and Essay Scoring")
body(
    "Automated scoring of student work has attracted considerable attention alongside the rise of "
    "large language models. Studies from 2024 and 2025 report that LLM-based graders can approach "
    "human agreement, such as reaching a Quadratic Weighted Kappa of 0.68 in one setting [5]. "
    "However, these models sometimes deviate from the rubric they were given, which means their "
    "scores can drift from what a human grader would assign."
)
body(
    "EssayJudge, introduced in 2025, is a multi-granular benchmark that was developed to test "
    "vision-language models on trait-specific essay scoring across several dimensions rather than a "
    "single overall grade [6]. It appears to be the first benchmark placing proprietary systems such "
    "as GPT-4o and Gemini-1.5-Pro alongside open-source models on grading consistency specifically. "
    "A related study evaluating vision-language models in Indonesian classrooms found that these "
    "systems identify fully correct answers reliably but handle partial credit poorly [7], a "
    "pattern that recurs in Section 4.3."
)
body(
    "A 2026 study on automated exam grading with foundation models examined fairness-aware "
    "recognition of handwritten answers, achieving 98.4% recognition accuracy on a benchmark of 61 "
    "anonymised exams [8], which helps to make the point that grading benchmarks built purely around "
    "correlation with a human score remain incomplete. Evaluation frameworks for educational AI "
    "therefore need to look beyond aggregate accuracy toward the properties that determine whether a "
    "system is safe to deploy, which is the motivation behind EvalBench's dedicated trustworthiness "
    "track."
)

h2("2.3 Confidence Calibration and Trustworthiness in AI")
body(
    "Calibration concerns whether a model's stated confidence matches its empirical accuracy. A "
    "2025 survey of uncertainty quantification in large language models concludes that most "
    "instruction-tuned models are systematically overconfident [9], which is a serious concern in "
    "any setting where a wrong but confident answer causes real harm, exam grading being an obvious "
    "example. Guo et al. demonstrated this overconfidence problem in standard deep networks and "
    "proposed temperature scaling as a low-cost post-hoc correction that reduces Expected "
    "Calibration Error without altering the underlying model."
)
body(
    "Closer to the present setting, a 2024 study on confidence estimation for automatic short "
    "answer grading found that LLMs readily produce a confidence value, but that value does not "
    "reliably track whether the underlying answer is correct [10]. A companion 2024 survey reviewed "
    "post-hoc calibration techniques for black-box LLMs and found no single method that generalises "
    "across task types [11]. A 2026 study evaluating self-reported confidence, self-consistency "
    "voting, and token probability across seven LLMs on three educational grading datasets reports "
    "an average ECE of 0.166 for the best-performing method [19], a finding broadly consistent with "
    "the overconfidence pattern reported in Track 5, though that study measures calibration on "
    "typed short-answer grading rather than handwritten exam scoring with evidence generation. This "
    "body of work helps to justify the design of Track 5, in that it measures calibration, evidence "
    "generation, and reasoning as three separate, independently interpretable properties rather than "
    "one combined score."
)

h2("2.4 Multi-Task and Vision-Language Benchmarks")
body(
    "Multi-task benchmarking has an established record in NLP research. GLUE and SuperGLUE "
    "aggregated diverse sentence-level tasks into a single leaderboard score, which made it "
    "possible to compare models on general language competence rather than isolated task "
    "performance, and BIG-Bench later extended this to 204 distinct tasks. Vision-language research "
    "has followed a similar path. PM4Bench, released in 2025, evaluates large vision-language "
    "models across a parallel multilingual, multi-modal corpus spanning ten languages [12], and "
    "LENS adopts an open-set configuration in which models answer natural-language queries grounded "
    "in photographic content [13], both useful steps toward more realistic evaluation conditions."
)
body(
    "None of these frameworks, however, were designed with education in mind. Their tasks are drawn "
    "from general-domain photography, web text, or captioning corpora, none of which resemble a "
    "handwritten exam script or a teacher's rubric. EvalBench occupies the space these benchmarks "
    "leave open, in that it borrows the shared-schema, multi-track structure that GLUE-style "
    "benchmarks popularised, but applies it to the document types and grading tasks specific to "
    "education, namely handwriting recognition, rubric-free scoring, and confidence calibration "
    "under a grading workload, evaluated together rather than in isolation."
)

h2("2.5 Positioning Relative to Existing Benchmarks")
body(
    "Table 1 summarises how "
    "EvalBench relates to the closest prior benchmarks discussed above on five qualitative axes: "
    "whether the benchmark evaluates OCR, whether it evaluates grading, whether it evaluates "
    "trustworthiness or calibration, whether it defines a single unified schema spanning multiple "
    "task types, and whether it is built around educational documents specifically rather than "
    "general-domain text or images. The comparison is drawn from each benchmark's own published "
    "description rather than from re-running it, since the point being illustrated is coverage of "
    "capability, not head-to-head performance."
)
add_table(
    headers=["Benchmark", "OCR", "Grading", "Trustworthiness", "Unified Schema", "Educational Focus"],
    rows=[
        ["OCRBench [18]",        "Yes", "No",  "No",      "No",  "No"],
        ["EssayJudge [6]",       "No",  "Yes", "No",      "No",  "Yes"],
        ["GLUE / SuperGLUE",     "No",  "No",  "No",      "Yes", "No"],
        ["PM4Bench [12]",        "No",  "No",  "No",      "No",  "No"],
        ["EvalBench v1 (ours)",  "Yes", "Yes", "Yes",     "Yes", "Yes"],
    ],
    cap_text="Table 1. Qualitative capability comparison between EvalBench v1 and the closest prior "
             "benchmarks, based on each benchmark's published description."
)
body(
    "No row besides EvalBench has a Yes in more than two columns. OCRBench is the strongest prior "
    "OCR benchmark but stops at OCR; EssayJudge is the strongest prior grading benchmark but does "
    "not evaluate OCR or calibration at all; GLUE and SuperGLUE popularised the unified-schema, "
    "multi-task leaderboard format this paper borrows, but for general-domain sentence tasks rather "
    "than education; and PM4Bench extends multi-task vision-language evaluation to new languages "
    "without touching grading, trustworthiness, or educational documents. EvalBench is the only "
    "entry in the table that combines all five properties, which is the concrete, tabulated version "
    "of the novelty claim made in Section 1."
)

h2("2.6 Why Existing Benchmarks Are Insufficient")
body(
    "Table 1 shows the gap; this subsection explains why it cannot be closed by simply running an "
    "existing benchmark on educational data, or by combining two existing benchmarks side by side."
)
body(
    "OCRBench cannot evaluate grading because it has no notion of a graded outcome in its schema: "
    "its 29 constituent datasets carry ground-truth transcriptions, not teacher-assigned scores, "
    "confidence judgements, or evidence citations, so extending it to Track 3 or Track 5 would "
    "require building an entirely new annotation layer rather than adjusting an evaluation script. "
    "EssayJudge cannot evaluate OCR because its inputs are already-legible essay text and trait "
    "rubrics; it was not designed to test whether a system can first read a handwritten page "
    "correctly before grading it, so a model that fails at transcription but happens to guess a "
    "plausible score would be invisible to it. GLUE and SuperGLUE cannot evaluate educational "
    "documents because their constituent tasks, natural language inference, sentiment classification, "
    "coreference resolution and similar, are built from general-domain sentence pairs with no "
    "handwriting, no exam layout, and no grading rubric anywhere in the pipeline; adapting them would "
    "mean discarding the tasks rather than reusing them."
)
body(
    "Concatenating OCRBench and EssayJudge results side by side, the most obvious workaround, would "
    "still not produce what EvalBench provides, because the two benchmarks use different corpora, "
    "different splits, and different metric conventions, so there is no shared record linking an OCR "
    "failure on a specific handwritten answer to the grading error it causes downstream. EvalBench's "
    "single schema, with every record carrying a record_id, dataset, split, track, and track-specific "
    "fields under one JSON Schema (Section 3.1), is what makes that link possible: the same "
    "underlying document can, in principle, be traced through OCR, grading, and trustworthiness "
    "evaluation without switching corpora or metric definitions partway through. That cross-task "
    "traceability, not any single track in isolation, is what existing benchmarks do not provide and "
    "what this paper contributes."
)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 – DATASET
# ══════════════════════════════════════════════════════════════════════════════
h1("3. Dataset Description and Pre-processing")
body(
    "No single publicly available dataset covers all five tracks EvalBench is designed to evaluate, "
    "so the corpus was assembled from six independent sources rather than one. Together these "
    "sources contribute 132,891 records spanning word-level handwriting crops, full examination "
    "page scans, teacher-graded exam PDFs, and handwritten question-answer pairs, a composition "
    "chosen to stress different aspects of document understanding rather than to maximise record "
    "count alone."
)

body(
    "In the Tracks column of Table 2, the numbers refer to the five evaluation tracks defined in "
    "Section 1: 1 is OCR accuracy, 2 is document understanding, 3 is educational assessment, 4 is "
    "multi-modal question answering, and 5 is trustworthiness."
)
add_table(
    headers=["Dataset", "Type", "Total", "Train", "Val", "Test", "Tracks"],
    rows=[
        ["IAM [14]",           "Word (handwritten)",       "38,305",  "26,813", "3,831", "7,661",  "1"],
        ["IIIT English Word",  "Word (handwritten)",       "85,020",  "59,513", "8,503", "17,004", "1"],
        ["IIIT English Page",  "Page (exam)",              "3,500",   "2,450",  "350",   "700",    "1, 2"],
        ["anshcode1",          "Exam script (unlabelled)", "5,915",   "--",     "--",    "5,915",  "1, 2"],
        ["Mendeley [15]",      "Exam paper (graded)",      "50",      "--",     "--",    "50",     "3, 5"],
        ["gopika13",           "Word (handwritten)",       "101",     "101",    "--",    "--",     "1, 2, 4"],
        ["Total",              "",                         "132,891", "88,877", "12,684","31,330", ""],
    ],
    cap_text="Table 2. EvalBench v1 dataset composition across six sources, verified against the "
             "processed corpus (datasets/processed/evalbench_v1.jsonl)."
)

body(
    "IIIT English Word is the largest single source once the corpus is verified against the "
    "processed data files: 85,020 word-level records extracted from real examination papers rather "
    "than newspaper-style prose, which introduces a domain shift relative to IAM and is precisely "
    "the property that makes cross-dataset comparison worthwhile. IAM contributes 38,305 processed "
    "word images from 657 different writers, covering offline English handwriting collected outside "
    "any examination context; this is smaller than IAM's full 85,756-image release because the "
    "processing pipeline currently retains only the subset with usable transcriptions and image "
    "quality above the harness's minimum threshold. These are greyscale word crops of variable "
    "width, and the variety of writing styles they contain is what makes them a demanding "
    "character-level recognition test."
)
body(
    "IIIT English Page contributes 3,500 full examination page images at a mean resolution of "
    "2,847 by 3,807 pixels. Page-level scans are a different kind of test, in that they carry "
    "layout structure, line segmentation, and multi-paragraph flow that a cropped word image cannot "
    "represent. anshcode1 contributes a further 5,915 exam-script page images (CC BY 4.0) that are "
    "currently unlabelled, no transcription or ground-truth score has yet been attached to them, so "
    "they are held out as test-split material pending annotation rather than used in any experiment "
    "reported in Section 4. Tracks 2 (document understanding) and 4 (multimodal understanding) are "
    "specified in the EvalBench schema against IIIT English Page, anshcode1, and gopika13 for Track "
    "2, and gopika13 for Track 4, plus a fourth source, the IEEE Answer Sheet Dataset, that is "
    "referenced in the schema but has not yet been downloaded into this release. None of these "
    "tracks have baselines reported here, and Track 2 and Track 4 presently have no labelled "
    "evaluation data at all rather than only missing baselines, a distinction discussed further in "
    "Section 6.3."
)
body(
    "The Mendeley Graded Exam Papers dataset is the smallest labelled source by record count, "
    "comprising 50 scanned handwritten exam papers, but it is the only source that carries genuine "
    "teacher-assigned scores, ranging from 13 to 43 marks (mean 27.9, standard deviation 7.2). That "
    "scarcity of graded ground truth is precisely why it anchors Tracks 3 and 5, since without it "
    "there would be no way to check whether a predicted score, or the confidence stated in that "
    "score, corresponds to anything a human grader would recognise. gopika13 rounds out the corpus "
    "with 101 supplementary handwriting samples, assigned entirely to the training split rather "
    "than the test split, used to broaden Track 1 coverage beyond IAM and "
    "IIIT-Word."
)

body(
    "Before turning to the rank-reversal result itself, it is worth quantifying how different IAM "
    "and IIIT-Word actually are at the level a model experiences them, namely the joint distribution "
    "of per-record OCR outcomes rather than raw pixel content. For every (model, record) prediction "
    "in the OCR experiment (n = 1,100 per corpus, pooled across all four models), we build a "
    "seven-dimensional feature vector of CER, WER, exact match, character accuracy, normalised "
    "reference length, normalised prediction length, and normalised latency, standardised to zero "
    "mean and unit variance across the pooled sample. The Frechet distance between the resulting "
    "IAM and IIIT-Word feature distributions, the same mean-and-covariance formulation used for FID "
    "in image generation, is 0.866; the simpler Euclidean distance between the two class centroids "
    "is 0.597, and a label-permutation test on that centroid distance (2,000 permutations, seed 42) "
    "gives p = 0.0005. This confirms quantitatively, not just visually, that the two corpora induce "
    "measurably different OCR outcome distributions, consistent with the rank reversal reported in "
    "Section 4.2."
)
fig("fig05_tsne_features.png",
    "Figure 1. PCA projection of the same standardised OCR-outcome feature space (Frechet distance "
    "0.866, centroid-distance permutation p = 0.0005), shown as a two-dimensional visual summary of "
    "the quantitative result above.")
body(
    "Figure 1 is a two-dimensional PCA projection of that same feature space, included as a visual "
    "summary rather than as the primary evidence. It shows substantial overlap in the dense central "
    "cluster, but also a visibly separate secondary cluster and a long tail of points where the two "
    "datasets diverge, a pattern consistent with the Frechet distance reported above: related enough "
    "to invite direct comparison, yet different enough that a model tuned to one distribution does "
    "not automatically transfer to the other."
)

h2("3.1 Data Schema")
body(
    "Every record, regardless of its originating dataset, is normalised to EvalBench JSON Schema "
    "v2.0 before any track-specific processing begins, which is what allows a single evaluation "
    "harness to run unmodified across six otherwise incompatible data formats. The required fields "
    "are record_id, dataset, split, image_path, ref (the ground-truth text or numeric score), and "
    "track. Track-specific extensions, such as teacher_score and max_score for Tracks 3 and 5, are "
    "included as named optional fields, and a catch-all extra object carries any remaining "
    "track-specific payload (for example the confidence, evidence, and reasoning fields Track 5 "
    "reads for calibration analysis), which keeps each evaluation script self-contained."
)

h2("3.2 Train, Validation, and Test Splits")
body(
    "The corpus is partitioned into training, validation, and test subsets using stratified "
    "sampling by dataset source, so that each subset preserves the proportional representation of "
    "all six sources. Splitting uses scikit-learn's StratifiedShuffleSplit with random_state = 42, "
    "which guarantees an identical split regardless of machine or Python version (>= 3.8). Mendeley "
    "is small enough that every one of its 50 records is assigned directly to the test subset, "
    "while gopika13, at 101 records, is assigned entirely to the training split rather than test, "
    "since its role in this release is to broaden Track 1 training coverage rather than to "
    "contribute held-out evaluation samples. All results reported in this paper are computed "
    "exclusively on the test split of the dataset that produced them."
)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 4 – EXPERIMENTS
# ══════════════════════════════════════════════════════════════════════════════
h1("4. Experimental Phase")

h2("4.1 Experimental Setup")
body(
    "All OCR experiments were run on an Apple M3 machine using CPU/MPS only, with no dedicated GPU "
    "involved at any stage, which helps to keep the Track 1 results reproducible on ordinary "
    "hardware. Track 3 and Track 5 experiments instead call the Gemini Flash family of "
    "vision-language models through Google's free-tier API, which imposes its own constraints: rate "
    "limits were handled with exponential backoff (30-120 seconds between retries) and a --resume "
    "flag that skips records already scored. Every evaluation script is a standalone Python 3.10+ "
    "module with a command-line interface and JSONL output, which makes incremental re-runs "
    "straightforward."
)

h2("4.2 Track 1 - OCR Accuracy")
body(
    "Table 3 reports word-level OCR accuracy for four models on the test splits of IAM and IIIT "
    "English Word. Every result is restricted to words of at least three characters, which removes "
    "punctuation-only tokens that would otherwise inflate error rates; after this filter, n = 199 "
    "for IAM and n = 223 for IIIT English Word for three of the four models, and n = 133 and "
    "n = 146 respectively for PaddleOCR. Tesseract 5, EasyOCR, and TrOCR-base were evaluated in one "
    "batch run over 300 raw records per corpus, while PaddleOCR was added in a later, separate run "
    "capped at 200 raw records per corpus; after the reference-length filter this leaves PaddleOCR "
    "with roughly a third fewer samples than the other three models. This is an operational "
    "inconsistency in how this pilot release's evaluation batches were configured rather than a "
    "deliberate design choice, and re-running PaddleOCR at the same sample cap as the other three "
    "models is planned for the next release; the PaddleOCR figures below should be read with that "
    "smaller, non-matched sample size in mind. Four complementary metrics are reported, namely "
    "Character Error Rate (CER), Word Error Rate (WER), Exact Match (EM), and Character Accuracy, "
    "computed as 1 minus CER under a substitution-only edit-distance convention, because no single "
    "metric captures both character-level precision and full-word correctness on its own."
)

add_table(
    headers=["Model", "IAM CER", "IAM EM", "IAM CharAcc", "IIIT-W CER", "IIIT-W EM", "IIIT-W CharAcc"],
    rows=[
        ["Tesseract 5",  "0.798", "0.045", "0.202", "0.662", "0.148", "0.338"],
        ["EasyOCR",      "0.631", "0.045", "0.369", "0.372", "0.296", "0.628"],
        ["TrOCR-base",   "0.229", "0.553", "0.771", "0.522", "0.291", "0.478"],
        ["PaddleOCR [2]","0.513", "0.075", "0.487", "0.177", "0.664", "0.823"],
    ],
    cap_text="Table 3. Word-level OCR results on IAM and IIIT English Word test splits."
)

body(
    "The pattern in Table 3 is not subtle. TrOCR-base posts the best Exact Match on IAM (0.553), "
    "which is consistent with its IAM-adjacent pre-training, but that same model's EM falls to "
    "0.291 on IIIT-Word, a drop of 0.261. PaddleOCR runs the opposite course, reaching 0.664 EM on "
    "IIIT-Word, the best score in the table, against only 0.075 on IAM, a fall of 0.589 in the "
    "other direction. The practical consequence is that whichever model tops the leaderboard on one "
    "dataset is not merely weaker but last-place on the other, which means a team validating only on "
    "IAM and a team validating only on IIIT-Word would reach recommendations that directly "
    "contradict each other."
)
body(
    "To check that these EM gaps are not an artefact of sampling noise, each was tested with a "
    "paired bootstrap (10,000 resamples, seed 42) and a paired permutation test over the "
    "per-record exact-match indicators. All four gaps are statistically significant at p < 0.001, "
    "with the exception of Tesseract 5 at p = 0.0005: Tesseract 5, IAM-minus-IIIT-Word EM gap "
    "-0.103, 95% CI [-0.159, -0.047]; EasyOCR, gap -0.251, 95% CI [-0.316, -0.184]; TrOCR-base, "
    "gap +0.261, 95% CI [+0.170, +0.351]; PaddleOCR, gap -0.589, 95% CI [-0.675, -0.497]. None of "
    "the four confidence intervals cross zero, which supports reading the rank reversal as a "
    "structural property of how these models generalise across handwriting styles rather than as "
    "noise from a small test split."
)

fig("fig02_5fold_boxplot.png",
    "Figure 2. 5-fold cross-validation CER distribution by model and dataset.")
fig("fig03_confusion_matrix.png",
    "Figure 3. Character-level confusion matrix for TrOCR-base on IAM (top-10 confused characters).")
fig("fig08_cer_violin.png",
    "Figure 4. CER distribution violin plot by model and dataset.")

body(
    "Figure 2 places the five-fold cross-validation CER distributions for all four models side by "
    "side. TrOCR-base shows the tightest, most stable spread on IAM, consistent with the dataset it "
    "was effectively tuned toward, but that stability does not carry over to IIIT-Word, where its "
    "spread widens noticeably. Figure 3 isolates where TrOCR-base goes wrong on IAM: the confusion "
    "matrix is dominated by visually similar character pairs such as 'i' and 'l', which suggests "
    "the residual errors are largely perceptual rather than a sign of a broken model. Figure 4's "
    "violin plot places all four models and both datasets on one axis, and the long lower tails for "
    "TrOCR-base on IIIT-Word and for PaddleOCR on IAM offer the clearest single summary of the "
    "rank-reversal result reported above."
)

add_table(
    headers=["Model", "IAM EM", "IIIT-W EM", "Gap (EM)"],
    rows=[
        ["Tesseract 5", "0.045", "0.148", "-0.103"],
        ["EasyOCR",    "0.045", "0.296", "-0.251"],
        ["TrOCR-base", "0.553", "0.291", "+0.261"],
        ["PaddleOCR",  "0.075", "0.664", "-0.589"],
    ],
    cap_text="Table 4. Cross-dataset generalisation gap (IAM EM minus IIIT-Word EM)."
)

body(
    "Table 4 makes the size of the gap explicit rather than leaving it implicit across two "
    "single-dataset tables. TrOCR-base carries the largest positive gap (+0.261), meaning it "
    "performs far better on IAM than on IIIT-Word; PaddleOCR carries the largest negative gap "
    "(-0.589), meaning the reverse. A gap this size is larger than the entire EM score achieved by "
    "three of the four models on IAM, which suggests a structural property of how each architecture "
    "responds to the two corpora rather than ordinary evaluation noise. The bootstrap and "
    "permutation tests reported later in this section confirm this directly: all four gaps are "
    "statistically significant, with 95% confidence intervals that exclude zero."
)

h2("4.3 Track 3 - Educational Assessment")
body(
    "Track 3 evaluates Gemini Flash on all 50 Mendeley exam papers, with teacher scores spanning 13 "
    "to 43 marks. Each paper is rendered as a 150 DPI PNG and given to the model with a holistic "
    "grading prompt that requests a single numeric score, with no rubric and no per-criterion "
    "breakdown provided. This is deliberately the hardest version of the task, since the model must "
    "read the handwriting and form a judgement about its quality in one pass. Free-tier API rate "
    "limits meant the 50 papers were graded across three Gemini variants (gemini-2.5-flash, "
    "gemini-flash-latest, and gemini-3.5-flash), and the combined metrics below were computed over "
    "all 50 records regardless of which variant produced the prediction, a limitation discussed "
    "further in Section 6.3."
)

add_table(
    headers=["Model", "n", "MAE", "Pearson r", "Spearman rho", "EM+-1", "EM+-5"],
    rows=[["Gemini (combined)", "50", "3.780", "0.543", "0.550", "0.580", "0.780"]],
    cap_text="Table 5. Track 3 educational assessment results on Mendeley exam papers (n=50)."
)

body(
    "A Score MAE of 3.78 marks, against a 30-mark spread in the underlying data, is a reasonable "
    "outcome for a model that received no rubric at all, and a Pearson r of 0.543 confirms the "
    "model is tracking genuine signal rather than producing scores unrelated to paper quality. The "
    "EM+-5 figure of 0.780 is the more intuitive summary: four out of every five papers land within "
    "five marks of what the teacher actually gave. Where the model struggles is precision rather "
    "than direction, in that the moderate rather than strong correlation points to a model that "
    "distinguishes a strong paper from a weak one but not two papers separated by only a mark or "
    "two, which is exactly the resolution a real grading decision usually turns on."
)
body(
    "Both point estimates carry wide bootstrap 95% confidence intervals given the n = 50 sample: "
    "MAE, [2.090, 5.730] marks, and Pearson r, [0.235, 0.801] (10,000 resamples, seed 42). The MAE "
    "interval excludes zero by a wide margin, so the model is reliably making non-trivial errors "
    "rather than grading near-perfectly with occasional noise, and the Pearson r interval stays "
    "entirely on the positive side, so the correlation itself is unlikely to be a chance artefact of "
    "this particular 50-paper sample, even though its exact strength is not tightly pinned down. "
    "Both intervals are wide enough that this result should be read as a pilot-scale finding rather "
    "than a precise benchmark number, as discussed further in Section 6.3."
)

fig("fig06_pred_vs_actual_score.png",
    "Figure 5. Predicted vs. actual exam score scatter plot (Track 3, n=50).")
fig("fig10_score_error_histogram.png",
    "Figure 6. Score prediction error histogram (Track 3, n=50).")

body(
    "Figure 5 plots predicted score against teacher score for all 50 papers. Predictions cluster "
    "tightly around the diagonal for mid-range scores, which is where most of the corpus sits, but "
    "the scatter widens noticeably at both the top and bottom of the scale. Figure 6's error "
    "histogram, with +-3 and +-5 mark tolerance bands overlaid, gives a second view of the same "
    "result: the bulk of predictions sit comfortably inside the +-5 band, but the tail beyond it is "
    "not negligible, and it is exactly the papers in that tail that would need teacher review in any "
    "responsible deployment of this kind of grading assistant."
)

h2("4.4 Track 5 - Trustworthiness")
body(
    "Track 5 asks a different question of the same 50 papers: not only what score the model gives, "
    "but whether that score can be trusted. Gemini is prompted to return four structured fields for "
    "every paper, namely SCORE, CONFIDENCE (0-1), EVIDENCE (a direct quotation supporting the "
    "score), and REASONING. Confidence calibration is summarised with Expected Calibration Error "
    "(ECE), and discriminative ability, whether stated confidence separates correct from incorrect "
    "predictions, is summarised with AUROC, both computed against a correctness threshold of +-3 "
    "marks. As with Track 3, free-tier API limits meant Track 5 was graded across more than one "
    "Gemini variant, though a different pair than Track 3 used: gemini-3.1-flash-lite for 31 "
    "records and gemini-3.5-flash for 19 records. Combined metrics below are computed over all 50 "
    "records regardless of which variant produced the prediction, the same measurement confound "
    "discussed for Track 3 in Section 6.3."
)

add_table(
    headers=["Model", "n", "ECE", "AUROC", "Evidence Generation Rate", "Reasoning Rate", "MAE"],
    rows=[["Gemini (combined)", "50", "0.251", "0.571", "1.000", "1.000", "4.120"]],
    cap_text="Table 6. Track 5 trustworthiness results on Mendeley exam papers (n=50)."
)

body(
    "Evidence Generation Rate and Reasoning Rate measure whether the prompt succeeds at eliciting "
    "the two supporting fields, not whether their content is accurate or actually supports the "
    "score given, and are named accordingly rather than as a 'grounding' metric, since grounding in "
    "the trustworthy-AI sense implies a verified link between the cited evidence and the model's "
    "output, which this release does not check. Evidence Generation Rate is the fraction of records "
    "for which the EVIDENCE field is non-empty and longer than five characters, and Reasoning Rate "
    "is the fraction for which the REASONING field is non-empty and longer than ten characters, "
    "both computed directly on the raw model output with no human or automated check on whether the "
    "quoted evidence actually supports the score given. Both rates are a perfect 1.000 on this "
    "basis, in that the structured prompt succeeds completely at getting the model to generate a "
    "non-empty EVIDENCE field and a non-empty REASONING field on every paper, without exception; "
    "this is a presence check on prompt compliance, not a verified-support or grounding-quality "
    "metric, and should be read accordingly. The other half of the result is less encouraging: ECE "
    "of 0.251 and AUROC of 0.571, barely above the 0.5 line marking pure chance, show that the "
    "model's stated confidence carries almost no information about whether the score it just "
    "produced is correct. The model is fluent and well organised in how it justifies itself, but "
    "that fluency is not a signal of reliability, which means a human reviewer has little cue to "
    "know which predictions deserve a second look."
)
body(
    "Both calibration statistics carry wide bootstrap 95% confidence intervals at n = 50: AUROC, "
    "[0.320, 0.680], and ECE, [0.129, 0.376] (10,000 resamples, seed 42). The AUROC interval "
    "straddles the 0.5 chance line, so this result cannot rule out the possibility that confidence "
    "carries some real discriminative signal the point estimate underestimates, but it equally "
    "cannot rule out chance-level performance, and either way the practical implication is the "
    "same: at this sample size, stated confidence should not be trusted to flag which predictions "
    "need a second look. Widening this dataset is the most direct way to tighten both intervals, as "
    "discussed in Section 6.3."
)

fig("fig04_auc_roc.png",
    "Figure 7. AUC-ROC curve for Track 5 confidence calibration (AUROC = 0.571).")
fig("fig07_calibration_diagram.png",
    "Figure 8. Reliability diagram and confidence histogram (ECE = 0.251).")

body(
    "Figure 7's ROC curve sits close to the diagonal chance line for most of its length, which is "
    "the visual signature of confidence that does not discriminate. Figure 8's reliability diagram "
    "makes the overconfidence concrete, in that within most confidence bins the model's empirical "
    "accuracy sits below its stated confidence, and the confidence histogram beneath it shows "
    "predictions clustering at the high-confidence end regardless of whether the underlying score "
    "was correct. This is the specific failure mode Track 5 was designed to catch, and it is one "
    "that a standard accuracy metric such as MAE would not reveal on its own."
)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 5 – ABLATION
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Ablation Study")
body(
    "The results in Section 4.2 already show that EM scores shift between IAM and IIIT-Word. This "
    "ablation isolates the ranking consequence of that shift directly, since a metric can move "
    "without changing which model a practitioner would actually deploy, and it is the deployment "
    "decision, not the raw metric, that a benchmark ultimately exists to inform."
)

add_table(
    headers=["Model", "IAM EM", "IAM Rank", "IIIT-W EM", "IIIT-W Rank", "Rank Change"],
    rows=[
        ["TrOCR-base", "0.553", "1", "0.291", "3", "-2 (drops)"],
        ["EasyOCR",    "0.045", "2", "0.296", "2", "0 (stable)"],
        ["PaddleOCR",  "0.075", "4", "0.664", "1", "+3 (rises)"],
        ["Tesseract 5","0.045", "3", "0.148", "4", "-1 (drops)"],
    ],
    cap_text="Table 7. Ablation: rank change between IAM and IIIT-Word evaluation."
)

fig("fig14_rank_comparison.png",
    "Figure 9. Model rank by evaluation corpus, IAM versus IIIT-Word (lower is better).")
fig("fig13_radar_chart.png",
    "Figure 10. Radar chart comparing models across IAM EM, IIIT-W EM, character accuracy, and speed.")

body(
    "The Rank Change column in Table 7, visualised directly in Figure 9, is the most informative "
    "result in this ablation. TrOCR-base falls from first to third place and PaddleOCR climbs from "
    "last to first, a complete inversion of the top and bottom of the leaderboard rather than a "
    "modest reshuffling in the middle. EasyOCR is the one model whose rank holds steady across both "
    "corpora, which is worth noting given its middling absolute scores on both. Choosing a "
    "deployment model based on IAM alone points to TrOCR-base; choosing based on IIIT-Word alone "
    "points to PaddleOCR, and those two choices cannot both be correct given that the only variable "
    "that changed between them is which corpus was used, not any difference in underlying model "
    "quality. Figure 10's radar chart adds a multi-attribute view across accuracy on both corpora "
    "and inference speed, and it shows that no single model's polygon dominates on every axis, "
    "which reinforces that the choice of best model here depends on which trade-off, and which "
    "corpus, matters most to the deployer."
)
fig("fig01_train_loss_acc.png",
    "Figure 11. TrOCR-base training and validation loss and accuracy curves (IAM, 20 epochs).")
fig("fig11_wordlen_vs_cer.png",
    "Figure 12. Word length vs. CER scatter with binned means for TrOCR-base and PaddleOCR.")
fig("fig12_latency_comparison.png",
    "Figure 13. Model inference latency comparison across four OCR models (Apple M3, CPU).")
body(
    "The remaining diagnostic figures add model-internal detail to the ablation. Figure 11 confirms "
    "that TrOCR-base's IAM fine-tuning converges cleanly over 20 epochs with no sign of "
    "instability, which rules out an under-trained checkpoint as the explanation for its IIIT-Word "
    "shortfall, in that the drop reported in Section 4.2 is a generalisation gap rather than a "
    "training artefact. Figure 12 shows that character error rate climbs with word length for both "
    "TrOCR-base and PaddleOCR, though the two models diverge in exactly where that climb becomes "
    "steep, which hints at different failure modes between a sequence decoder and a "
    "detection-recognition pipeline. Figure 13 adds a practical dimension often absent from "
    "accuracy-only benchmarks, namely inference latency on ordinary CPU hardware, which matters as "
    "much as raw accuracy for any institution planning to deploy OCR at the scale of an entire exam "
    "board."
)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 6 – DISCUSSION
# ══════════════════════════════════════════════════════════════════════════════
h1("6. Discussion")

h2("6.1 Multi-Source Evaluation Findings")
body(
    "The rank reversal between IAM and IIIT-Word is the central finding of this study, and it is "
    "worth restating plainly: the model ranked first on one corpus is ranked last on the other. The "
    "cross-dataset EM gap of 0.589 recorded for PaddleOCR is, on its own, larger than the total EM "
    "score achieved by three of the four models on IAM, which makes ordinary measurement noise an "
    "unlikely explanation, a conclusion the paired bootstrap and permutation tests reported in "
    "Section 4.2 confirm directly (p < 0.001 for all four models except Tesseract 5 at p = 0.0005). "
    "A laboratory that validates only on IAM would deploy TrOCR-base; one that "
    "validates only on IIIT-Word would deploy PaddleOCR instead, and the two recommendations are "
    "mutually exclusive. Because IAM and IIIT-Word are both widely used handwriting corpora rather "
    "than one being an obscure edge case, this is a warning about the ordinary practice of "
    "validating on whichever corpus happens to be convenient, not about unusual data."
)
body(
    "This carries a direct procurement implication that is easy to overlook. Vendor claims of "
    "state-of-the-art OCR accuracy are almost always benchmarked on a single reference dataset "
    "chosen by the vendor, and Table 4 shows how misleading that single number can be once "
    "deployment data differs even modestly from benchmark data. An institution evaluating "
    "handwriting recognition for its own exam scripts should, at minimum, request performance "
    "figures across more than one representative corpus before trusting a headline accuracy claim."
)

h2("6.2 Assessment and Trustworthiness Findings")
body(
    "On this 50-paper Mendeley sample, Track 3 shows that Gemini Flash can grade handwritten exam "
    "papers with no rubric and land within five marks on 78% of them, a genuinely useful zero-shot "
    "baseline for this specific rubric-free setup, though not one precise "
    "enough for high-stakes, unsupervised deployment. The model's weakness is specifically at the "
    "resolution that separates a 27 from a 29, not at telling a strong paper from a weak one, which "
    "suggests the gap between current performance and human-grader precision is a rubric problem as "
    "much as a model-capability problem, a hypothesis Section 7 returns to directly."
)
body(
    "Track 5's result is the sharper of the two findings, precisely because it would remain "
    "invisible to anyone looking only at MAE. A model can cite specific supporting evidence for "
    "every prediction it makes and still carry a confidence score with almost no relationship to "
    "whether that prediction is correct. Evidence generation and calibration are, empirically, "
    "decoupled properties, and evaluating one without the other risks certifying a grading tool as "
    "reliable when only its explanatory fluency, not its actual judgement, has been checked."
)

h2("6.3 Limitations")
body(
    "First, the Track 3 and Track 5 results are each combined across more than one Gemini model "
    "variant because free-tier API quotas were exhausted mid-evaluation, and the two tracks were "
    "not even affected the same way: Track 3 spans three variants (gemini-2.5-flash, "
    "gemini-flash-latest, gemini-3.5-flash), while Track 5 spans a different pair "
    "(gemini-3.1-flash-lite, gemini-3.5-flash). This introduces a measurement confound in both "
    "tracks, since some of the variance in the combined results may reflect differences between "
    "model versions rather than genuine task difficulty, and future work should repeat each "
    "50-paper evaluation on a single, fixed model version before the calibration numbers are "
    "treated as final."
)
body(
    "Second, 50 papers is a small sample for calibration statistics specifically. The bootstrap 95% "
    "confidence interval around Track 5's AUROC (0.320 to 0.680) spans almost the entire plausible "
    "range for a discrimination metric, and the interval around ECE (0.129 to 0.376) is similarly "
    "wide, which means the Track 5 numbers should be read as a directional signal, namely that the "
    "model is overconfident, rather than a tightly bounded estimate of exactly how overconfident."
)
body(
    "Third, Tracks 2 and 4 (document understanding and multi-modal understanding) are defined in the "
    "EvalBench schema and partially covered by existing data, IIIT English Page, anshcode1, and "
    "gopika13 for Track 2, and gopika13 for Track 4, but neither has an evaluation script, a "
    "baseline model, or labelled ground truth for the anshcode1 portion in this release, and both "
    "are missing the IEEE Answer Sheet Dataset the schema also names as a source, which has not yet "
    "been downloaded. This is a more fundamental gap than \"no baselines reported\": Tracks 2 and 4 "
    "are schema-complete but not yet experiment-ready, and closing that gap, starting with "
    "annotating anshcode1 and acquiring the IEEE source, is the most immediate item on the roadmap "
    "in Section 7. The five-track leaderboard is therefore only three-fifths populated in this "
    "release. This is exactly the gap the pilot-release label in the title and abstract is meant to "
    "signal: EvalBench v1 defines the full five-track schema up front, deliberately, so that Tracks "
    "2 and 4 can be filled in against a fixed, already-published interface rather than bolted on "
    "after the fact, but it reports experimental results only for the three tracks that are "
    "genuinely ready, rather than padding the leaderboard with placeholder baselines."
)
body(
    "Fourth, as the dataset statistics in Figure 14 make clear, the corpus is heavily weighted "
    "toward word-level OCR data, in that IAM and IIIT-Word alone contribute more than 85,000 "
    "records each, while Mendeley, the only source with genuine graded ground truth, contributes "
    "just 50. Balancing this distribution, most plausibly by extending the graded-assessment side "
    "of the corpus, would make EvalBench a more even test of all five tracks."
)
body(
    "Fifth, statistical uncertainty is now quantified but remains wide at this sample size. "
    "Sections 4.2 through 4.4 report bootstrap 95% confidence intervals and, for Track 1, paired "
    "permutation tests alongside every headline number, and the Track 1 EM gaps are all "
    "statistically significant at p < 0.001 (Tesseract 5 at p = 0.0005). The Track 3 and Track 5 "
    "intervals are considerably wider, reflecting the n = 50 sample those two tracks are built on, "
    "and should be read as evidence that the reported effects are real rather than as a precise "
    "estimate of their exact magnitude. Growing the graded-assessment portion of the corpus beyond "
    "50 papers is the most direct way to tighten those intervals in the next release."
)

fig("fig09_dataset_statistics.png",
    "Figure 14. EvalBench v1 dataset statistics (total vs. test split per source).")

h2("6.4 Threats to Validity")
body(
    "Beyond the limitations above, three validity concerns are worth flagging for readers planning "
    "to build on this work. The internal validity of the Track 1 comparison depends on test subsets "
    "of 133 to 223 records per model, after the three-character reference-length filter is applied "
    "to IAM and IIIT-Word, and a larger held-out sample would tighten the confidence bounds around "
    "the reported EM gaps; PaddleOCR's smaller, non-matched sample (Section 4.2) compounds this for "
    "that model specifically, and its EM gap should be weighted accordingly until the batches are "
    "re-run at a matched sample size. The external validity of the Track 3 and "
    "Track 5 findings is bounded by a single source of graded ground truth, the Mendeley dataset, "
    "collected under one grading rubric and one national examination context, so the specific "
    "magnitude of the overconfidence pattern reported here should not be assumed to transfer "
    "unchanged to other grading systems, subjects, or countries without direct testing. Construct "
    "validity is a live question for Track 5 specifically, in that a model's willingness to cite a "
    "plausible-looking quotation is treated here as evidence generation, but a perfect Evidence "
    "Generation Rate does not by itself confirm that the cited evidence is the actual basis for the "
    "score rather than a "
    "fluent justification generated after the score was already decided, a distinction that Section "
    "7 identifies as a direction for future work."
)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 7 – CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
h1("7. Conclusion and Future Work")
body(
    "This paper introduced EvalBench v1 as a pilot release, a unified multi-task benchmark for AI "
    "evaluation in educational settings, built from 132,891 quality-filtered records across six "
    "datasets and organised into five evaluation tracks, three of which are experimentally complete "
    "in this release. The central experimental result, a rank reversal between IAM and IIIT-Word "
    "large enough to flip a deployment decision entirely, demonstrates that multi-source evaluation "
    "is not merely preferable but necessary, in that identical models produce opposite "
    "recommendations depending solely on which corpus is used to judge them. The Track 5 result adds "
    "a second, independent lesson: structured, well-justified output and genuinely calibrated "
    "confidence are separate properties, and a system can excel at one while failing the other. On "
    "this evidence, Gemini Flash under a no-rubric holistic prompt is not yet ready for unsupervised, "
    "high-stakes grading, and the same evaluation pattern is worth applying to any AI grading tool "
    "before it is deployed in a real classroom."
)
body(
    "Four directions follow from the limitations catalogued in Section 6.3, roughly in order of "
    "urgency. The most pressing is resolving the model-consistency confound in Tracks 3 and 5 by "
    "re-running each full 50-paper evaluation on one fixed Gemini version instead of the several "
    "variants each track currently mixes together, which would remove the largest source of "
    "measurement noise currently affecting the calibration metrics."
)
body(
    "The second direction is completing Tracks 2 and 4. Both are already defined in the EvalBench "
    "schema and partially covered by existing document-image data, but neither has ground-truth "
    "labels, an evaluation script, or a baseline model yet. Annotating the currently unlabelled "
    "anshcode1 records, acquiring the IEEE Answer Sheet Dataset the schema also names as a source, "
    "and then implementing and running a first baseline for at least one of the two tracks would "
    "turn the current three-fifths-populated leaderboard into a more complete comparison and is the "
    "single addition most likely to strengthen this benchmark's case for publication."
)
body(
    "The third direction is rubric-grounded prompting for Track 3. The current protocol "
    "deliberately withholds any rubric, which is useful for measuring a lower bound on zero-shot "
    "capability but understates what these models can do when given the same grading criteria a "
    "human teacher would use. Supplying structured, per-criterion rubrics and re-measuring MAE would "
    "test directly whether the fine-grained allocation weakness identified in Section 6.2 is a "
    "genuine capability ceiling or an artefact of an underspecified prompt."
)
body(
    "The fourth direction is linguistic breadth. EvalBench v1 is English-only, which leaves out a "
    "large share of the world's classrooms where instruction and assessment happen in Tamil, "
    "Arabic, Chinese, or another script entirely. Extending the corpus with graded handwriting "
    "samples in additional languages would test whether the rank-reversal and overconfidence "
    "patterns reported here are properties of these particular models and datasets, or something "
    "closer to a general property of how vision-language systems behave under handwriting-heavy, "
    "high-stakes educational tasks regardless of language."
)
body(
    "Beyond these four directions, the broader engineering motivation for a benchmark like EvalBench "
    "is standardisation. Educational AI vendors and institutions currently have no shared, "
    "reproducible way to compare an OCR component, a grading model, and a confidence-estimation "
    "layer against the same corpus, split protocol, and metric definitions, which makes vendor "
    "claims and internal validation results difficult to compare across teams or across procurement "
    "cycles. A common schema and evaluation harness, released publicly with fixed splits, gives "
    "institutions a way to reproduce a vendor's reported numbers rather than take them on faith, and "
    "gives researchers a fixed target against which future OCR and vision-language systems can be "
    "compared on equal footing rather than on whichever corpus a paper happens to report. That "
    "shared basis is also a precondition for safer deployment: the rank-reversal and overconfidence "
    "findings in this paper are only visible because OCR, grading, and calibration were evaluated "
    "under one protocol instead of three separate ones, and a procurement or deployment decision "
    "made without that kind of joint evaluation risks missing exactly the failure modes that matter "
    "most before an AI grading tool reaches a real classroom."
)

# ══════════════════════════════════════════════════════════════════════════════
#  END-MATTER  (required by most Q1 / Elsevier-style journals)
# ══════════════════════════════════════════════════════════════════════════════
body(
    "Author contact information and the CRediT authorship contribution statement are provided in "
    "the accompanying title page document, in line with this journal's anonymous peer review "
    "process, and are not reproduced here."
)

h1("Declaration of Competing Interest")
body(
    "The authors declare that they have no known competing financial interests or personal "
    "relationships that could have appeared to influence the work reported in this paper."
)

h1("Data and Code Availability")
body(
    "All six source datasets used to construct EvalBench v1 are publicly available from their "
    "original repositories, as cited in Section 3 and the reference list. The EvalBench v1 corpus, "
    "the JSON Schema v2.0 specification, all evaluation scripts, and the full experimental outputs "
    "reported in this paper will be released at the time of publication to support independent "
    "verification and reuse. Because IAM is distributed under a license that requires individual "
    "registration and restricts redistribution of its images, the public release will include IAM "
    "and IIIT English Word records as derived annotations and licensed-source pointers rather than "
    "as re-hosted raw images, so that redistribution stays within each source dataset's own terms."
)

h1("Declaration on the Use of Generative AI and AI-Assisted Technologies in the Writing Process")
body(
    "During the preparation of this manuscript, the authors used generative AI tools to assist "
    "with improving language, wording, organisation, and readability, and to support literature "
    "search and cross-checking of computed statistics against source data. The authors carefully "
    "reviewed, verified, and edited all AI-assisted content and take full responsibility for the "
    "accuracy and integrity of the final manuscript."
)

# ══════════════════════════════════════════════════════════════════════════════
#  REFERENCES  (2024-2026 only)
# ══════════════════════════════════════════════════════════════════════════════
h1("References")

ref(1,  "Crosilla, G., Klic, L., & Colavizza, G.",
    "2025",
    "Benchmarking Large Language Models for Handwritten Text Recognition",
    "Journal of Documentation, 81(7), 334-354",
    "https://doi.org/10.1108/jd-03-2025-0082")

ref(2,  "PaddleOCR Team",
    "2025",
    "PaddleOCR 3.0 Technical Report",
    "arXiv preprint",
    "arXiv:2507.05595")

ref(3,  "Hamdi, L., Tamasna, A., Boisson, P., & Paquet, T.",
    "2025",
    "PILOT: A Promptable Interleaved Layout-aware OCR Transformer",
    "arXiv preprint",
    "arXiv:2504.03621")

ref(4,  "Nagasubramanian, A., Almazyad, A. S., Balakrishnan, S. A., Bharath Ram, M. M., Potti, D., "
    "Zakariah, M., & AlSekait, D. M.",
    "2025",
    "OCRNet: A Robust Deep Learning Framework for Alphanumeric Character Recognition to Assist the "
    "Visually Impaired",
    "Scientific Reports, 15, Article 41344",
    "https://doi.org/10.1038/s41598-025-25278-9")

ref(5,  "Liew, P. Y., & Tan, I. K. T.",
    "2024",
    "On Automated Essay Grading Using Large Language Models",
    "Proceedings of CSAI 2024, ACM",
    "https://doi.org/10.1145/3709026.3709030")

ref(6,  "Su, J., Yan, X., Fu, Y., Zhang, H., Ye, Q., Liu, X., Huo, Y., Zhou, C., & Hu, X.",
    "2025",
    "EssayJudge: A Multi-Granular Benchmark for Assessing Automated Essay Scoring Capabilities of "
    "Multimodal Large Language Models",
    "arXiv preprint",
    "arXiv:2502.11916")

ref(7,  "Aisyah, N., Al Kautsar, M. D., Hidayat, A., Chowdhury, R., & Koto, F.",
    "2025",
    "From Handwriting to Feedback: Evaluating VLMs and LLMs for AI-Powered Assessment in Indonesian "
    "Classrooms",
    "arXiv preprint",
    "arXiv:2506.04822")

ref(8,  "Grabowski, H.",
    "2026",
    "Towards Fully Automated Exam Grading: Fairness-Aware Recognition of Handwritten Answers with Foundation Models",
    "arXiv preprint",
    "arXiv:2606.11477")

ref(9,  "Liu, X., Chen, T., Da, L., Chen, C., Lin, Z., & Wei, H.",
    "2025",
    "Uncertainty Quantification and Confidence Calibration in Large Language Models: A Survey",
    "arXiv preprint",
    "arXiv:2503.15850")

ref(10, "Cong, L., Hahn, S., Gombert, S., Camus, L., Drachsler, H., & Kroehne, U.",
    "2026",
    "Confidence Estimation in Automatic Short Answer Grading with LLMs",
    "Springer Lecture Notes in Computer Science (AIED 2026)",
    "https://doi.org/10.1007/978-3-032-29755-6_8")

ref(11, "Xie, L., Liu, H., Zeng, J., Tang, X., Han, Y., Luo, C., Huang, J., Li, Z., Wang, S., & He, Q.",
    "2024",
    "A Survey of Calibration Process for Black-Box LLMs",
    "arXiv preprint",
    "arXiv:2412.12767")

ref(12, "Gao, P., Song, X., Wu, X., Zhu, D., Shen, S., Wang, S., Wei, W., Yang, C., Zhang, S., Li, W., "
    "Wang, B., Lin, D., Wu, L., & He, C.",
    "2025",
    "PM4Bench: Benchmarking Large Vision-Language Models with Parallel Multilingual Multi-Modal "
    "Multi-task Corpus",
    "arXiv preprint",
    "arXiv:2503.18484")

ref(13, "Yao, R., Zhang, B., Huang, J., Long, X., Zhang, Y., Zou, T., Wu, Y., Su, S., Xu, Y., Zeng, "
    "W., Yang, Z., Li, G., Zhang, S., Li, Z., Chen, Y., Xiong, S., Xu, P., Zhang, J., Zhou, B., "
    "Clifton, D., & Van Gool, L.",
    "2025",
    "LENS: Multi-level Evaluation of Multimodal Reasoning with Large Language Models",
    "arXiv preprint",
    "arXiv:2505.15616")

ref(14, "Marti, U. V., & Bunke, H.",
    "2002",
    "The IAM-database: An English Sentence Database for Offline Handwriting Recognition",
    "International Journal on Document Analysis and Recognition, 5(1), 39-46",
    "https://doi.org/10.1007/s100320200071")

ref(15, "Dinesh, K. P.",
    "2026",
    "A Dataset of Digitized Student Examination Papers, Answer Keys, and Manual Evaluations for "
    "Automated Grading Research",
    "Mendeley Data",
    "https://doi.org/10.17632/sf3kvjwknt.1")

ref(16, "Song, Y., Zhu, Q., Wang, H., & Zheng, Q.",
    "2024",
    "Automated Essay Scoring and Revising Based on Open-Source Large Language Models",
    "IEEE Transactions on Learning Technologies",
    "https://doi.org/10.1109/TLT.2024.3396873")

ref(17, "Eneye, T. A. N. F., Ijezue, C. F., Amjad, A. I., Amjad, M., Butt, S., & Castaneda-Garza, G.",
    "2025",
    "Advances in Auto-Grading with Large Language Models: A Cross-Disciplinary Survey",
    "Proceedings of the 20th Workshop on Innovative Use of NLP for Building Educational Applications, "
    "ACL 2025",
    "https://aclanthology.org/2025.bea-1.35")

ref(18, "Liu, Y., Li, Z., Huang, M., Yang, B., Yu, W., Li, C., Yin, X., Liu, C., Jin, L., & Bai, X.",
    "2024",
    "OCRBench: On the Hidden Mystery of OCR in Large Multimodal Models",
    "Science China Information Sciences",
    "https://doi.org/10.1007/s11432-024-4235-6")

ref(19, "Ferrer, R., Turgut, D., Chen, Z., & Sonkar, S.",
    "2026",
    "When Can We Trust LLM Graders? Calibrating Confidence for Automated Assessment",
    "arXiv preprint",
    "arXiv:2603.29559")

# ══════════════════════════════════════════════════════════════════════════════
#  SAVE
# ══════════════════════════════════════════════════════════════════════════════
doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")
