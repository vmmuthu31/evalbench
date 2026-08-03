"""
Build the EvalBench Title Page (with author details) as a .docx file.
Submitted alongside the blinded manuscript for Engineering Applications of
Artificial Intelligence (Elsevier), which uses anonymous peer review.
Run:  python3 evalbench_paper2/paper_draft/build_title_page.py
"""

from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT     = Path(__file__).parent.parent.parent
OUT_PATH = ROOT / "evalbench_paper2/paper_draft/EvalBench_TitlePage.docx"

doc = Document()

section = doc.sections[0]
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin   = Inches(1.2)
section.right_margin  = Inches(1.2)


def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.size = Pt(13)
    p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    return p


def body(text, indent=True):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.3)
    for run in p.runs:
        run.font.size = Pt(11)
    return p


# ── Title ────────────────────────────────────────────────────────────────────
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run(
    "EvalBench v1 (Pilot Release): A Unified Multi-Task Benchmark for AI Evaluation in "
    "Educational Settings"
)
run.bold = True
run.font.size = Pt(16)
doc.add_paragraph()

# ── Authors and affiliations ────────────────────────────────────────────────
h1("Authors")
body(
    "Vairamuthu M.¹ and Sudhan M. B.¹",
    indent=False,
)
body(
    "¹ Department of Computer Science and Engineering, SRM Institute of Science and "
    "Technology, Chennai, India.",
    indent=False,
)

h1("Corresponding Author")
body(
    "Vairamuthu M., vm8470@srmist.edu.in, Department of Computer Science and Engineering, "
    "SRM Institute of Science and Technology, Chennai, India.",
    indent=False,
)

h1("All Author Emails")
body("Vairamuthu M.: vm8470@srmist.edu.in", indent=False)
body("Sudhan M. B.: sudhanm@srmist.edu.in", indent=False)

# ── CRediT ───────────────────────────────────────────────────────────────────
h1("CRediT Authorship Contribution Statement")
body(
    "Vairamuthu M.: Conceptualization, Methodology, Software, Data Curation, Formal Analysis, "
    "Investigation, Visualization, Writing - Original Draft."
)
body(
    "Sudhan M. B.: Conceptualization, Supervision, Validation, Writing - Review and Editing, "
    "Project Administration."
)

# ── Funding ──────────────────────────────────────────────────────────────────
h1("Funding")
body(
    "This research did not receive any specific grant from funding agencies in the public, "
    "commercial, or not-for-profit sectors."
)

# ── Acknowledgements ─────────────────────────────────────────────────────────
h1("Acknowledgements")
body(
    "The authors thank the maintainers of the IAM, IIIT English Word/Page, anshcode1, "
    "Mendeley Graded Exam Papers, and gopika13 datasets for making their data publicly "
    "available, which made this benchmark possible."
)

# ── Declaration of Competing Interest (repeated for convenience) ───────────
h1("Declaration of Competing Interest")
body(
    "The authors declare that they have no known competing financial interests or personal "
    "relationships that could have appeared to influence the work reported in this paper."
)

doc.save(str(OUT_PATH))
print(f"Saved: {OUT_PATH}")
